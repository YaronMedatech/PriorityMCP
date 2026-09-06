# Build the Hebrew system-administrator guide as a .docx.
#
# Hebrew forces the shape of this: python-docx over reportlab, because reportlab
# has no bidi shaping and would render Hebrew reversed unless a separate shaper
# and a Hebrew TTF were wired in. Word already does both, and exports PDF in one
# step, so the document is authored once and converted by whoever needs it.
#
# Every paragraph is explicitly marked RTL. Word does not infer direction from
# content, so a paragraph left unmarked renders left-aligned with the punctuation
# in the wrong place -- readable, and obviously produced by a machine.

import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = sys.argv[1]

doc = Document()

# Base style: a font that actually carries Hebrew glyphs, and one that carries
# them in the East-Asian slot too, or Word substitutes silently.
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)
style_fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
style_fonts.set(qn("w:eastAsia"), "Arial")
style_fonts.set(qn("w:cs"), "Arial")


def rtl(paragraph):
    """Mark a paragraph and every run inside it right-to-left."""
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:bidi"))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in paragraph.runs:
        rPr = run._element.get_or_add_rPr()
        rPr.append(OxmlElement("w:rtl"))
        rPr.get_or_add_rFonts().set(qn("w:cs"), "Arial")
    return paragraph


def h(text, level):
    p = doc.add_heading(text, level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        run.font.name = "Arial"
    return rtl(p)


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return rtl(p)


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return rtl(p)


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return rtl(p)


def code(text):
    """Code stays left-to-right. A path or a JSON block mirrored is unusable."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Consolas")
    fonts.set(qn("w:hAnsi"), "Consolas")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    shade(p, "F2F2F2")
    return p


def shade(paragraph, hex_colour):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_colour)
    pPr.append(shd)


def note(text, colour="FFF4CE"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = False
    shade(p, colour)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return rtl(p)


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.RIGHT
    # Mirror the column order for RTL, or the header sits over the wrong column.
    tblPr = t._tbl.tblPr
    tblPr.append(OxmlElement("w:bidiVisual"))
    for i, text in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        rtl(p)
    for row in rows:
        cells = t.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.add_run(text)
            rtl(p)
    doc.add_paragraph()
    return t


# ---------------------------------------------------------------------------
title = doc.add_heading("PriorityMCP — מדריך מנהל מערכת", 0)
rtl(title)
para("חיבור מודלי שפה למערכת Priority ERP דרך שרת MCP").italic = True
para("עודכן: ספטמבר 2026")
doc.add_page_break()

h("1. מה זה, בשתי פסקאות", 1)
para(
    "PriorityMCP הוא שרת שמתווך בין מודל שפה (Claude, Gemini) לבין התקנת Priority ERP חיה. "
    "הוא חושף למודל שלושה־עשר כלים: עשרה לקריאה — חיפוש מסכים, קריאת עמודות, שאילתות, "
    "סיכומים — ושלושה להרצת תוכניות ודוחות של Priority."
)
para(
    "העיקרון שמנחה אותו הוא שהמודל שואל את Priority מה הדברים אומרים, במקום להסיק משמות "
    "המסכים. שמות המסכים הם קודים אנגליים אטומים שמשמעותם חיה בכותרת העברית, והסקה מהאחד "
    "על השני היא בדיוק הטעות שהשרת קיים כדי למנוע."
)

h("2. הארכיטקטורה בפועל", 1)
table(
    ["רכיב", "היכן", "תפקיד"],
    [
        ["שרת PriorityMCP", "TATINT (192.168.50.12)", "שירות Windows, מאזין ב-HTTPS על פורט 3401"],
        ["Priority ERP", "tattest.tat.local", "מקור הנתונים, נקרא ב-OData ובערוץ Web SDK"],
        ["מכונת קצה", "הלפטופ של כל משתמש", "מריצה Claude Desktop / Claude Code / אחר"],
    ],
)
para(
    "השרת קורא מ-Priority בשני ערוצים נפרדים, וההפרדה נכפית ע\"י Priority ולא נבחרה: "
    "OData אינו יכול להריץ תוכנית, וה-Web SDK דורש כתובת וזהות משלו."
)

h("3. דרישות מוקדמות לשרת", 1)
bullet("Windows Server עם Node.js גרסה 20 ומעלה")
bullet("הרשאות מנהל מקומי — לרישום השירות ולפתיחת פורט בחומת האש")
bullet("גישה רשתית ל-Priority בפורט 443")
bullet("משתמש Priority שמורשה למסכי מחולל: EFORM, EXEC, EREP, EPROG")

note(
    "ההרשאה האחרונה קריטית ולא מובנת מאליה. בלי EFORM לא ניתן לבנות את מילון המסכים, "
    "וכל כלי הגילוי נופלים — המודל חוזר לנחש שמות מסכים, וזה הכשל שהמערכת נועדה למנוע. "
    "המסכים האלה עונים HTTP 400 כשהם סגורים, ואותה הודעה בדיוק מופיעה גם כשהמשתמש חסר הרשאה."
)

h("4. התקנת שרת חדש", 1)
para("כל השרשרת מתבצעת בסקריפט אחד. ב-PowerShell כמנהל, מתוך תיקיית הפרויקט:")
code("powershell -ExecutionPolicy Bypass -File scripts\\bootstrap-server.ps1")
para("הסקריפט מבצע, בסדר הזה, ונעצר בכשל ראשון:")
numbered("בדיקות מקדימות: הרשאות מנהל, קובץ .env, גרסת Node, תלויות")
numbered("הנפקת תעודת TLS לשם המכונה הזו ולכתובות ה-IP שלה")
numbered("סיבוב סיסמת התעודה וטוקן הגישה, וכתיבתם ל-.env")
numbered("התאמת נתיב Node בהגדרת השירות")
numbered("בדיקת טיפוסים, שש־עשרה חבילות בדיקה, ובדיקה חיה מול Priority")
numbered("בניית מילון המסכים — זהו השער האמיתי, ובלעדיו אין טעם להמשיך")
numbered("רישום השירות ופתיחת פורט 3401 בחומת האש")
numbered("בדיקת בריאות ויצירת חבילת התקנה ללקוחות")
note(
    "הסקריפט בטוח להרצה חוזרת: תעודה שכבר נושאת את שם המכונה נשמרת, והטוקן מסובב פעם אחת "
    "בלבד — כדי שניסיון שני לא יבטל הגדרות שכבר חולקו למשתמשים.",
    "E8F4EA",
)

h("5. הוספת משתמש — קודם כל, איזה מודל זהות", 1)
para(
    "זה קובע אם יש בכלל שלב בשרת. בדוק לפני כל דבר אחר:"
)
code("Test-Path C:\\PriorityMCP\\users.json")
table(
    ["", "לא קיים — זהות משותפת", "קיים — זהות לכל קורא"],
    [
        ["שלב בשרת", "אין", "שורה בקובץ users.json"],
        ["מה בקובץ של המשתמש", "שם משתמש וסיסמת Priority", "טוקן אקראי, בלי סיסמה"],
        ["מי Priority חושב ששואל", "אותו משתמש", "אותו משתמש"],
        ["ביטול גישה לאדם אחד", "החלפת סיסמה ב-Priority", "מחיקת שורה והפעלה מחדש"],
    ],
)
para(
    "שני המודלים שומרים על הרשאות Priority אישיות. ההבדל הוא היכן הסיסמה נמצאת: על כל לפטופ, "
    "או פעם אחת על השרת.",
    bold=True,
)
note(
    "מומלץ לעבור לזהות לכל קורא. סיסמה בקובץ הגדרות מועתקת, מגובה ונשלחת במייל; טוקן ניתן "
    "לביטול בפני עצמו ואינו שווה דבר במקום אחר. המעבר הוא העתקת users.example.json ל-users.json "
    "והפעלה מחדש."
)

h("6. צד השרת — הוספת רשומה", 1)
para("רלוונטי רק במודל זהות לכל קורא. במודל המשותף אין מה לעשות בשרת.")
h("6.1 ייצור טוקן", 2)
code("[Convert]::ToBase64String((1..32 | ForEach-Object { Get-Random -Max 256 }))")
para("טוקן קצר מ-16 תווים נדחה בעלייה עם אזהרה, ולא מתקבל בשקט.")
h("6.2 הוספה לקובץ", 2)
code(
    '{\n'
    '  "users": [\n'
    '    {\n'
    '      "token": "<הטוקן מסעיף 6.1>",\n'
    '      "label": "sivan-laptop",\n'
    '      "priorityUser": "<שם המשתמש ב-Priority>",\n'
    '      "priorityPass": "<הסיסמה>"\n'
    '    }\n'
    '  ]\n'
    '}'
)
para(
    "השדה label מופיע בלוג השרת ליד כל קריאה של אותו אדם, ולכן לעולם אינו סוד. "
    "אם ל-Priority יש PAT, עדיף להחליף את שני שדות הסיסמה ב-priorityToken."
)
h("6.3 הפעלה מחדש ואימות", 2)
code("Restart-Service PriorityMCP\nGet-Content C:\\PriorityMCP\\service\\logs\\priority-mcp.err.log -Tail 20")
para(
    "בלוג מופיעה שורה עם מספר הזהויות שנטענו. רשומה עם טוקן קצר, כפול או חסר אישורים "
    "מדולגת עם אזהרה — כלומר ספירה נמוכה מהצפוי אומרת שאחת נדחתה, והאזהרה אומרת איזו."
)
note("הקובץ users.json מוחרג מ-git מפני שהוא מחזיק אישורים. ודא ב-git status אם ערכת אותו בתוך תיקיית הפרויקט.")

h("7. צד המשתמש — שלושה שלבים, בסדר הזה", 1)
para(
    "הסדר הוא העיקר. כל שלב ניתן לאימות בנפרד, כך שכשמשהו נכשל הוא מזהה את עצמו במקום "
    "שיחפשו אותו במקום הלא נכון.",
    bold=True,
)

h("7.1 אמון בתעודה", 2)
para(
    "כל הקליינטים האלה רצים על Node, ו-Node אינו קורא את מחסן התעודות של Windows. "
    "לכן ייבוא התעודה ל-Windows לא יעזור להם, וגם תעודה שהונפקה מה-CA של הדומיין לא תעזור. "
    "מה שעובד הוא משתנה הסביבה NODE_EXTRA_CA_CERTS.",
)
para("הדבק את הבלוק שמנהל המערכת מספק — הוא כותב את התעודה ומגדיר את המשתנה. אינו דורש הרשאות מנהל.")

h("7.2 אימות לפני הגדרה", 2)
code("curl --cacert C:\\priority-mcp\\mcp-ca.pem --ssl-revoke-best-effort https://192.168.50.12:3401/health")
para('מצופה: {"ok":true,"transport":"streamable-http","auth":"bearer"}')
table(
    ["מה שרואים", "מה זה אומר"],
    [
        ["שגיאת תעודה", "שלב 7.1 לא נקלט. בדוק את טביעת האצבע"],
        ["timeout", "ניתוב או חומת אש. השרת פתוח לפרופילי Domain ו-Private בלבד"],
        ["connection refused", "השירות מושבת. בדוק Get-Service PriorityMCP"],
    ],
)
note(
    "לעולם אל תוסיף ‎-k או ‎--insecure כדי לעבור את השלב הזה. הם מקבלים כל שרת שמתחזה לכתובת, "
    "וזה הדבר היחיד שהתעודה מונעת. הדגל ‎--ssl-revoke-best-effort אינו כזה: curl ב-Windows דורש "
    "מקור לבדיקת ביטול, ו-CA פרטי אינו מפרסם אחד. החתימה ושם המארח עדיין נבדקים.",
    "FBE9E7",
)

h("7.3 הגדרת הקליינט", 2)
para("בכל אחד מהקליינטים, בלוק ההזדהות הוא אחד משניים, לפי המודל מסעיף 5:")
code('"headers": { "X-Priority-User": "...", "X-Priority-Pass": "..." }')
code('"headers": { "Authorization": "Bearer <הטוקן>" }')

h("8. הגדרה לפי סוג הקליינט", 1)

h("8.1 Claude Code — מאומת", 2)
para("קובץ בשם .mcp.json בשורש התיקייה שהמשתמש פותח:")
code(
    '{\n'
    '  "mcpServers": {\n'
    '    "priority": {\n'
    '      "type": "http",\n'
    '      "url": "https://192.168.50.12:3401/mcp",\n'
    '      "headers": { "X-Priority-User": "...", "X-Priority-Pass": "..." }\n'
    '    }\n'
    '  }\n'
    '}'
)
para("לאחר מכן לסגור ולפתוח את VS Code או הטרמינל לגמרי — לא Reload — ולהריץ ‎/mcp‎.")
note(
    "מלכודת: Windows Explorer מסתיר סיומות מוכרות, ולכן שינוי שם של קובץ טקסט ל-.mcp.json "
    "מייצר בפועל .mcp.json.txt והקליינט לא רואה אותו. בדוק עם Get-ChildItem בתוספת ‎-Force‎.",
    "FBE9E7",
)

h("8.2 Claude Desktop — מאומת חלקית", 2)
para("Settings ← Developer ← Edit Config. לא Connectors; בגרסאות ישנות אין לשונית כזו.")
para(
    "הקובץ כבר מכיל תוכן. יש להוסיף את mcpServers לצדו ולא להדביק מעליו, אחרת המשתמש מאבד "
    "את ההעדפות שלו. שים לב לפסיק אחרי הסוגר של preferences — בלעדיו ה-JSON אינו תקין "
    "ו-Claude Desktop מתעלם מהקובץ כולו בשקט, כולל ההגדרות האישיות.",
    bold=True,
)
note(
    "לא אומת אם הגרסה הנוכחית מקבלת type: http בקובץ הזה. Claude Desktop הגדיר היסטורית "
    "שרתי stdio בלבד. אם השרת אינו מופיע לאחר הפעלה מחדש מלאה — זו התשובה, והחלופה היא "
    "גשר mcp-remote, שדורש Node על אותה מכונה."
)

h("8.3 VS Code Copilot", 2)
para(
    "יש להשתמש בפקודה ולא בכתיבת קובץ ביד: Command Palette ← MCP: Add Server ← HTTP ← הכתובת. "
    "VS Code כותב את הקובץ בעצמו, וכך נמנעת ניחוש סכימה שהשתנתה בין גרסאות."
)

h("8.4 Gemini CLI", 2)
para("בקובץ ‎~/.gemini/settings.json‎ — שים לב ל-httpUrl ולא url:")
code(
    '{\n'
    '  "mcpServers": {\n'
    '    "priority": {\n'
    '      "httpUrl": "https://192.168.50.12:3401/mcp",\n'
    '      "headers": { "X-Priority-User": "...", "X-Priority-Pass": "..." },\n'
    '      "timeout": 120000,\n'
    '      "trust": false,\n'
    '      "excludeTools": ["run_program"]\n'
    '    }\n'
    '  }\n'
    '}'
)

h("9. תחזוקה שוטפת", 1)
table(
    ["משימה", "פקודה"],
    [
        ["פריסת שינוי בקוד", "git pull ואז npm run typecheck ואז Restart-Service PriorityMCP"],
        ["בדיקת מצב השירות", "Get-Service PriorityMCP"],
        ["צפייה בלוג", "Get-Content service\\logs\\priority-mcp.err.log -Tail 40"],
        ["בדיקת בריאות", "curl ‎--cacert‎ ... https://<שרת>:3401/health"],
    ],
)
note(
    "הרצת npm run typecheck לפני הפעלה מחדש אינה טקס. השירות מנסה לעלות שלוש פעמים ואז "
    "נשאר מושבת, כך ששגיאת תחביר אחת היא השבתה.",
    "FBE9E7",
)

h("10. אבטחה — מה שחשוב שמנהל המערכת יידע", 1)
bullet(
    "אישורי גישה לעולם אינם ארגומנט של כלי. הם מגיעים ככותרות HTTP או דרך בקשת הזדהות של "
    "הקליינט, כדי שלא ייכנסו לשיחה, לתמלול או ללוג הקריאות."
)
bullet("כותרות עם אישורים נדחות על חיבור לא מוצפן. זו אינה הגדרה שניתן לכבות.")
bullet(
    "רשימת החברות ב-PRIORITY_ENVIRONMENTS היא רשימת היתר ולא רמז: השם נכנס לנתיב URL, "
    "וזה בדיוק המקום שדרכו נעשה בעבר מעבר לחברה שכנה."
)
bullet(
    "המשתנה MCP_AUTH_TOKEN מעניק את זהות השרת עצמה. יש להתייחס אליו כאישור ניהולי ולא "
    "כמשהו שמחלקים."
)
bullet(
    "PRIORITY_ALLOW_ALL_PROGRAMS=1 מאפשר הרצה של כל ‎9,200‎ הפרוצדורות והדוחות, כולל כאלה "
    "שמעדכנות ומוחקות. למשתמש שרק שואל שאלות עדיף להחריג את כלי ההרצה בהגדרה שלו, או "
    "להגדיר PRIORITY_READ_ONLY=1 בשרת."
)

h("11. הסרת משתמש", 1)
para(
    "בזהות לכל קורא: מחיקת השורה מ-users.json והפעלה מחדש. הטוקן מפסיק לעבוד מיידית "
    "ושום דבר אחר לא מושפע."
)
para(
    "בזהות משותפת: אין ביטול בצד השרת. הקובץ אצל האדם מחזיק את סיסמת ה-Priority שלו, "
    "ולכן התשובה היא החלפת סיסמה. באסימטריה הזו נמצא הטיעון החזק ביותר לעבור למודל "
    "הראשון לפני שמספר המשתמשים גדל.",
    bold=True,
)

doc.save(OUT)
print("written:", OUT)
